"""Verify the three JPS-compliance fixes applied to SUBMIT_FINAL."""
import io
import re
import sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

SUBMIT = r"D:\Project Python\PythonProject9\Paper 7\Paper4_UQ_Comparison\SUBMIT_FINAL_2026-05-18"

doc = Document(SUBMIT + r"\Paper4_Manuscript.docx")
full_text = "\n".join(p.text for p in doc.paragraphs)

print("=== Issue 3: data-availability placeholder ===")
if "<author>" in full_text:
    print("  FAIL: <author> still present")
else:
    if "github.com/trangtt/Paper4_UQ_Comparison" in full_text:
        print("  OK: github.com/trangtt/Paper4_UQ_Comparison present")

print()
print("=== Issue 1: trim verification ===")
checks = [
    ("Old phrase removed: We do not claim parity with relative-window", "We do not claim parity with relative-window", False),
    ("New phrase present: Code, predictions, and metrics are released for full reproducibility", "Code, predictions, and metrics are released for full reproducibility", True),
    ("Old phrase removed: and broader UQ benchmarks", "and broader UQ benchmarks", False),
    ("New phrase present: Adjacent UQ work does not provide multi-method", "does not provide multi-method UQ comparisons for forward early-cycle knee prediction", True),
    ("Old phrase removed: sharpness–coverage frontier (ncal = 22", "sharpness–coverage frontier (ncal = 22", False),
    ("New phrase present: §S2.7–S2.9 cover the sharpness", "§S2.7–S2.9 cover the sharpness", True),
    ("Old phrase removed: Physics or aggressive stabilisation; the latter recovers", "Physics or aggressive stabilisation; the latter recovers", False),
    ("Old phrase removed: §6.4 recommendation (DE + split-conformal + CQR head) maps", "§6.4 recommendation (DE + split-conformal + CQR head) maps", False),
    ("Old phrase removed: Headline benchmark: 110 Severson LFP cells. Cross-chemistry", "Headline benchmark: 110 Severson LFP cells. Cross-chemistry", False),
    ("Old phrase removed: Statistical claims are conditional on the cell-level pairing", "Statistical claims are conditional on the cell-level pairing", False),
]
all_pass = True
for desc, needle, expect_present in checks:
    present = needle in full_text
    ok = (present == expect_present)
    if not ok:
        all_pass = False
    print(f"  {'OK' if ok else 'FAIL'} [{desc}] -> present={present}")

print()
print("=== Word counts ===")
paras = [p.text for p in doc.paragraphs]
start = next((i for i, t in enumerate(paras) if t.strip().startswith("1. Introduction")), 0)
end = next(
    (i for i, t in enumerate(paras) if i > start and (t.strip().startswith("Acknowledgement") or t.strip().startswith("CRediT"))),
    len(paras),
)
main = "\n".join(paras[start:end])
main_words = len([w for w in re.split(r"\s+", main) if w])
print(f"  Manuscript main text: {main_words}  (JPS limit 8000)")

cdoc = Document(SUBMIT + r"\Paper4_CoverLetter.docx")
ctext = "\n".join(p.text for p in cdoc.paragraphs)
cwords = len([w for w in re.split(r"\s+", ctext) if w])
print(f"  Cover letter: {cwords}  (JPS target <1 page)")

print()
print("=== Cover letter compliance ===")
has_reviewer = bool(re.search(r"(?i)Howey|Braatz|Yarin|Cand[eè]s|reviewer", ctext))
has_funding = bool(re.search(r"(?i)funding|grant", ctext))
print(f"  Contains reviewer mentions? {has_reviewer}  (JPS: should be NO)")
print(f"  Contains funding mentions? {has_funding}  (JPS: should be NO)")
print(f"  Dated 18 May 2026? {'18 May 2026' in ctext}")
print(f"  Has Trang signature? {'Tran Thanh Trang' in ctext}")
print(f"  Has Tam signature? {'Tran Nhut Tam' in ctext}")
print(f"  Has title? {'Uncertainty Quantification for Lithium-Ion Battery Knee-Point' in ctext}")

print()
print("=== Cover letter content preview (first 800 chars) ===")
print(ctext[:800])

print()
print("=== ALL CHECKS PASS ===" if all_pass else "=== SOME CHECKS FAILED ===")
