"""Fix abstract reference style per JPS guide.

JPS Guide for Authors > Abstract:
  "Avoid references. If any are essential to include, ensure that you cite
   the author(s) and year(s)."

Paper 4 abstract had a numbered reference [62] (BatteryGPT).
Replace with author+year format "(Hu et al., 2025)".
"""
import io
import re
import sys
from pathlib import Path
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

MANUSCRIPT = Path(
    r"D:\Project Python\PythonProject9\Paper 7\Paper4_UQ_Comparison\SUBMIT_FINAL_2026-05-18\Paper4_Manuscript.docx"
)

OLD = "(BatteryGPT 13 cycles at 30 % of variable lifetime [62])"
NEW = "(BatteryGPT, Hu et al. 2025, 13 cycles at 30 % of variable lifetime)"


def replace_in_para(p, old, new) -> bool:
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new, 1)
    p.runs[0].text = new_full
    for r in p.runs[1:]:
        r.text = ""
    return True


doc = Document(MANUSCRIPT)

# Find Abstract section
in_abstract = False
fixed = False
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith("Abstract"):
        in_abstract = True
        continue
    if t.startswith("Keywords") or t.startswith("Nomenclature"):
        in_abstract = False
        break
    if in_abstract and replace_in_para(p, OLD, NEW):
        fixed = True
        break

if fixed:
    doc.save(MANUSCRIPT)
    print(f"  OK: replaced [62] in abstract -> 'Hu et al. 2025' format")
else:
    print(f"  FAIL: phrase '{OLD}' not found in abstract")

# Verify post-save
doc2 = Document(MANUSCRIPT)
abs_text = ""
in_abs = False
for p in doc2.paragraphs:
    t = p.text.strip()
    if t.startswith("Abstract"):
        in_abs = True
        continue
    if t.startswith("Keywords") or t.startswith("Nomenclature"):
        break
    if in_abs:
        abs_text += " " + p.text

bracket_refs = re.findall(r"\[\d+\]", abs_text)
print(f"  Abstract bracketed refs after fix: {bracket_refs}  (target: [])")
print(f"  Abstract word count: {len([w for w in abs_text.split() if w])}")
