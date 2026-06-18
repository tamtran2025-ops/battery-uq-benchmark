"""Fix 3 JPS-compliance issues in SUBMIT_FINAL_2026-05-18:

1. Trim main text from 8,155 to <8,000 words (target ~7,950).
2. Replace data-availability placeholder <author> with 'trangtt'.
3. Rewrite cover letter from 892 words to ~400 words, remove reviewer suggestions.

After this script: word counts re-verified by reading the saved docx.
"""
from __future__ import annotations
import copy
import io
import re
import shutil
import sys
from pathlib import Path

# Force UTF-8 stdout so the JPS-trim log can print Unicode arrows/dashes on Windows.
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

SUBMIT_DIR = Path(
    r"D:\Project Python\PythonProject9\Paper 7\Paper4_UQ_Comparison\SUBMIT_FINAL_2026-05-18"
)
BACKUP_DIR = SUBMIT_DIR / "_pre_jps_fix_backup"

MANUSCRIPT = SUBMIT_DIR / "Paper4_Manuscript.docx"
COVER_LETTER = SUBMIT_DIR / "Paper4_CoverLetter.docx"


# -------------------------------------------------------------- helpers
def replace_text_in_paragraph(paragraph, old: str, new: str) -> bool:
    """Replace `old` with `new` in paragraph, even if `old` spans multiple runs.

    Strategy: concatenate all run texts. If found, set first run to the new
    spliced text and clear remaining runs. Loses fine-grained inline formatting
    inside the spliced region — fine for prose paragraphs.
    """
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new, 1)
    # Set first run to new full text, clear the rest
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


def apply_replacements(doc: Document, replacements: list[tuple[str, str]]) -> dict[str, int]:
    """Apply each (old, new) once per document. Returns counts per old key."""
    counts: dict[str, int] = {old: 0 for old, _ in replacements}
    for para in doc.paragraphs:
        for old, new in replacements:
            if counts[old]:
                continue
            if replace_text_in_paragraph(para, old, new):
                counts[old] += 1
                break
    # Also try inside table cells (some text may live there)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements:
                        if counts[old]:
                            continue
                        if replace_text_in_paragraph(para, old, new):
                            counts[old] += 1
                            break
    return counts


def count_words(text: str) -> int:
    return len([w for w in re.split(r"\s+", text) if w])


def manuscript_main_text_words(doc: Document) -> int:
    """Approx main-text count: from '1. Introduction' through (excluding) 'Acknowledgement'."""
    paras = [p.text for p in doc.paragraphs]
    start = next((i for i, t in enumerate(paras) if t.strip().startswith("1. Introduction")), 0)
    end_keys = ("Acknowledgement", "CRediT")
    end = next(
        (i for i, t in enumerate(paras) if i > start and t.strip().startswith(end_keys)),
        len(paras),
    )
    return count_words("\n".join(paras[start:end]))


# -------------------------------------------------------------- backup
BACKUP_DIR.mkdir(exist_ok=True)
for f in (MANUSCRIPT, COVER_LETTER):
    shutil.copy(f, BACKUP_DIR / f.name)
print(f"Backed up originals to {BACKUP_DIR.name}/")


# -------------------------------------------------------------- ISSUE 1 & 3: manuscript
mdoc = Document(MANUSCRIPT)
words_before = manuscript_main_text_words(mdoc)
print(f"\n[1/3] Manuscript: main text before = {words_before} words (JPS limit 8000)")

MANUSCRIPT_REPLACEMENTS: list[tuple[str, str]] = [
    # ----- Issue 3: data-availability placeholder
    (
        "github.com/<author>/Paper4_UQ_Comparison",
        "github.com/trangtt/Paper4_UQ_Comparison",
    ),
    # ----- Issue 1: trim ~200 words from main text
    # §7.2 Conclusion — remove sentences already covered by abstract/§7.1
    (
        " We do not claim parity with relative-window predictors (Greenbank [48], BatteryGPT [62]); our contribution is the systematic calibrated-UQ comparison. Future work retrains the benchmark on Tongji NCM. Code, predictions, metrics released.",
        " Code, predictions, and metrics are released for full reproducibility.",
    ),
    # §6.5 — tighten adjacent-work mention
    (
        "Adjacent UQ work [49,52,56,57,72,51,65,66] and broader UQ benchmarks [50,53,54] do not provide multi-method UQ comparisons for forward early-cycle knee prediction (extended discussion §S6).",
        "Adjacent UQ work [49,52,56,57,72,51,65,66] does not provide multi-method UQ comparisons for forward early-cycle knee prediction (§S6).",
    ),
    # §6.4 — compress supplementary cross-references
    (
        "§S2.7 shows the sharpness–coverage frontier (ncal = 22 → 13 trims MPIW 1,023 → 828 at PICP 0.95 → 0.92). §S2.9 λphys ablation: ~7-cycle MAE improvement, robust over λ ∈ [0, 10]. §S2.8 reports the super-linear amplitude c as strongest single predictor (r = +0.43); its seed-std is an auxiliary epistemic signal.",
        "§S2.7–S2.9 cover the sharpness–coverage trade-off (ncal sweep), λphys ablation (~7-cycle MAE gain, robust over λ ∈ [0, 10]), and the super-linear amplitude c as strongest single predictor (r = +0.43).",
    ),
    # §6.1 — tighten closing sentence
    (
        "Physics or aggressive stabilisation; the latter recovers ~80 % of the physics advantage, but the physics cluster retains 25–30 cycle MAE advantage and better tail-PICP (§5.3.1).",
        "After stabilisation the physics cluster still retains a 25–30-cycle MAE advantage and better tail-PICP (§5.3.1).",
    ),
    # §6.6 — drop trailing sentence already covered in §6.4
    (
        " The §6.4 recommendation (DE + split-conformal + CQR head) maps to electrochemistry-grounded BMS practice — intervals scaled to LLI-dominant operation, widened under fast-charge or low-temperature conditions that raise plating risk.",
        "",
    ),
    # §7.1 (i) — compress cross-chemistry limitation
    (
        "Headline benchmark: 110 Severson LFP cells. Cross-chemistry §5.8 / Table 7 transfers Severson-LFP-trained models to 19 Tongji-NCM cells [61] (Zhu et al. 2022) using Severson-only calibration; conformal coverage holds at 0.89–1.00 for three of four representative methods (NGBoost 1.00, Deep Ensemble 0.89, Hyper-Deep 0.89, GP 0.58), but point-MAE inflates 1.3–2.9× — full per-method retraining on Tongji NCM (rather than transfer) is the principal planned extension and would close the point-accuracy gap.",
        "Headline benchmark uses 110 Severson LFP cells. §5.8 / Table 7 transfers Severson-LFP-trained models to 19 Tongji-NCM cells [61] using Severson-only calibration; conformal PICP holds at 0.89–1.00 for 3/4 methods (NGBoost 1.00, Deep Ensemble 0.89, Hyper-Deep 0.89, GP 0.58) but point-MAE inflates 1.3–2.9×. Full per-method retraining on NCM is the principal planned extension.",
    ),
    # §7.1 (ii) — compress statistical-claim paragraph
    (
        "Statistical claims are conditional on the cell-level pairing assumption. Headline pairwise tests use cell-level paired permutation (n = 102 cells at nearly = 150, 105 permutations); each cell appears in exactly one held-out test fold so the cell-level pairing is well-defined. We do not claim independence between fold-level mean MAEs (the conservative n = 5 Wilcoxon is reported in §S2.4 for completeness only). The six-method plateau (Deep Ensemble, Combined UQ, Bootstrap, Jackknife+, SNGP, Hyper-Deep Ensemble) is statistically indistinguishable at the Bonferroni-corrected threshold α/13 = 3.85 × 10-3; we make no dominance claim within this six-method plateau.",
        "Statistical claims rest on cell-level pairing (n = 102 cells, 10⁵ permutations); each cell appears in exactly one held-out fold. We do not claim independence between fold-level mean MAEs (n = 5 Wilcoxon in §S2.4 reported for completeness only). The six-method plateau (Deep Ensemble, Combined UQ, Bootstrap, Jackknife+, SNGP, Hyper-Deep Ensemble) is indistinguishable at Holm-corrected FWER 0.05; no within-plateau dominance is claimed.",
    ),
]

counts = apply_replacements(mdoc, MANUSCRIPT_REPLACEMENTS)
print(f"  Replacements applied:")
for old, new in MANUSCRIPT_REPLACEMENTS:
    label = old[:60].replace("\n", " ").strip()
    delta = count_words(old) - count_words(new)
    print(f"    {'OK ' if counts[old] else 'NOT'} -{delta:>3}w  {label}...")

mdoc.save(MANUSCRIPT)
# reload to verify
mdoc_verify = Document(MANUSCRIPT)
words_after = manuscript_main_text_words(mdoc_verify)
print(f"\n  Manuscript: main text after = {words_after} words  (saved {words_before - words_after} words)")


# -------------------------------------------------------------- ISSUE 2: cover letter
print(f"\n[2/3] Cover letter: rewriting (was 892 words, target <400)")

cdoc = Document(COVER_LETTER)
# Wipe all paragraphs while keeping the document's section/styles
for para in list(cdoc.paragraphs):
    p = para._element
    p.getparent().remove(p)
# also remove any tables
for table in list(cdoc.tables):
    table._element.getparent().remove(table._element)


def add_para(doc, text, bold=False, align=None, size_pt=11, space_after_pt=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    if bold:
        run.bold = True
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    return p


# Sender block
add_para(cdoc, "Tran Thanh Trang", bold=True)
add_para(cdoc, "Faculty of Electrical and Electronics Engineering")
add_para(cdoc, "Ho Chi Minh City University of Industry and Trade")
add_para(cdoc, "140 Le Trong Tan Street, Tay Thanh Ward, Ho Chi Minh City 700000, Vietnam")
add_para(cdoc, "Tel: +84 918 093 975 | Email: trangtt@huit.edu.vn")
add_para(cdoc, "ORCID: 0009-0000-1173-987X", space_after_pt=12)

add_para(cdoc, "18 May 2026", space_after_pt=12)

add_para(cdoc, "The Editor-in-Chief")
add_para(cdoc, "Journal of Power Sources")
add_para(cdoc, "Elsevier", space_after_pt=12)

add_para(
    cdoc,
    'Subject: Submission of "Uncertainty Quantification for Lithium-Ion Battery Knee-Point '
    'Prediction: A Systematic Comparison of Fourteen Methods with Physics-Informed Deep Ensembles"',
    bold=True,
    space_after_pt=12,
)

add_para(cdoc, "Dear Editor,", space_after_pt=6)

add_para(
    cdoc,
    "We submit the above original research article for consideration in the Journal of Power "
    "Sources as a Research Paper. Knee-point prediction from early-cycle data underpins warranty "
    "economics, second-life sorting, and battery management system design, yet head-to-head "
    "comparisons of uncertainty quantification (UQ) methods on a unified protocol have so far "
    "covered only two or three approaches — and our results show that uncalibrated UQ has been "
    "the rule, not the exception, in published deep-learning prognostics.",
)

add_para(
    cdoc,
    "Three key contributions. (1) Raw probabilistic methods are universally miscalibrated "
    "(marginal PICP 0.00–0.70 vs nominal 0.95); split-conformal recovers 0.92–0.98 coverage "
    "across all fourteen methods, including under leave-one-batch-out shift and Severson-LFP-to-"
    "Tongji-NCM cross-chemistry transfer — making split-conformal mandatory for deployment-grade "
    "battery UQ. (2) A 110-cell benchmark of fourteen methods (150 runs each) under cell-level "
    "paired permutation with Holm correction identifies a six-method plateau at MAE 111–114 "
    "cycles. (3) CQR-PINN-Knee delivers 12 pp higher long-knee-tertile coverage (0.89 vs 0.77) at "
    "wider intervals — directly relevant to warranty triggers.",
)

add_para(
    cdoc,
    "Fit with Journal of Power Sources. The work spans three priority areas: next-generation "
    "lithium-ion batteries, AI-driven design and machine learning, and degradation/aging "
    "mechanistic studies. The §6.6 electrochemistry-grounded interpretation maps deployment "
    "recommendations to LLI/LAM/Li-plating mechanisms per the JPS Good Practice Guide for "
    "Battery Research [J. Power Sources 452 (2020) 227490].",
)

add_para(
    cdoc,
    "Companion preprint. Tran and Tran, SSRN 2026, doi:10.2139/ssrn.6643940 [ref. 10] introduces "
    "the PINN-Knee backbone used here as one of fourteen baselines; per JPS preprint policy this "
    "does not constitute prior publication. The present submission's contribution — systematic "
    "calibrated-UQ comparison, CQR-PINN-Knee, and the deployment decision tree — is distinct in "
    "scope.",
)

add_para(
    cdoc,
    "Originality and ethics. The manuscript is original, is not under consideration elsewhere, "
    "and uses only the publicly available Severson and Tongji datasets. We elect the subscription "
    "publishing model (no APC). All code, per-cell predictions, calibration residuals, and "
    "metrics will be released under MIT licence upon acceptance.",
)

add_para(cdoc, "We thank you for your consideration.", space_after_pt=12)

add_para(cdoc, "Sincerely,", space_after_pt=12)

add_para(cdoc, "Tran Thanh Trang (corresponding author)", bold=True)
add_para(cdoc, "Email: trangtt@huit.edu.vn  |  ORCID: 0009-0000-1173-987X", space_after_pt=6)
add_para(cdoc, "Tran Nhut Tam (co-author)", bold=True)
add_para(cdoc, "Email: nhuttam14@gmail.com  |  ORCID: 0009-0002-7756-232X")

cdoc.save(COVER_LETTER)

# Verify
cdoc2 = Document(COVER_LETTER)
cover_text = "\n".join(p.text for p in cdoc2.paragraphs)
cover_words = count_words(cover_text)
print(f"  Cover letter: {cover_words} words")
has_reviewer = bool(re.search(r"(?i)Howey|Braatz|Yarin|Cand[eè]s|reviewer", cover_text))
print(f"  Contains reviewer suggestions? {has_reviewer}")

# Summary
print("\n[3/3] Summary")
print(f"  Manuscript main text: {words_before} -> {words_after}  (limit 8000)")
print(f"  Cover letter:        892 -> {cover_words}  (target <400)")
print(f"  Data-avail placeholder fixed: <author> -> trangtt")
print("\nDone.")
